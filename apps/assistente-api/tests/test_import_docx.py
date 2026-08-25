from io import BytesIO

from docx import Document

from app.services.llm.prompts import PROJECT_DOC_SECTIONS
from app.services.wizards.text_import import split_by_headings


def test_split_by_headings_legacy_flow_maps_to_methods():
    text = """Contexto
Intro.

Fluxo de Análise
Executar pipeline diário.
"""
    sections = split_by_headings(text)
    assert "Intro" in sections["background"]
    assert "pipeline" in sections["methods_analysis"]


def test_docx_text_extraction():
    doc = Document()
    doc.add_heading("Contexto", level=1)
    doc.add_paragraph("Texto de contexto do estudo.")
    doc.add_heading("Objetivos", level=1)
    doc.add_paragraph("Objetivo principal.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    loaded = Document(buffer)
    full_text = "\n".join(p.text for p in loaded.paragraphs if p.text.strip())
    sections = split_by_headings(full_text)
    assert "contexto" in sections["background"].lower() or "Texto" in sections["background"]
    assert len(PROJECT_DOC_SECTIONS) == 11
